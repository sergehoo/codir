"""Export d'un CR CODIR vers .docx — format Kaydan.

Génère un Word avec :
  - En-tête (logo placeholder, date, référence, présidence, rapporteur)
  - Section Présents / Absents
  - Ordre du jour
  - Tableau des décisions/actions par direction (mirror du PDF importé)
"""
from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from apps.common.enums import AttendanceStatus

if TYPE_CHECKING:
    from apps.meetings.models import Meeting


# ─── Mappings d'affichage ────────────────────────────────────────
_TASK_STATUS_LABEL = {
    "todo": "Non démarré",
    "in_progress": "En cours",
    "blocked": "En attente",
    "overdue": "En retard",
    "done": "Terminé",
    "cancelled": "Annulé",
}


def _shade(cell, hex_color: str):
    """Applique un fond de couleur à une cellule (XML manipulation)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_col_widths(table, widths_cm):
    """Définit les largeurs de colonnes (en cm)."""
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def build_codir_minutes_docx(meeting: "Meeting") -> BytesIO:
    """Construit le .docx à partir d'un Meeting Django et retourne un BytesIO."""
    from apps.action_plans.models import ActionTask
    from apps.decisions.models import Decision

    doc = Document()

    # ─── Marges (A4 paysage utilitaire mais on reste portrait ici) ──
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # ─── En-tête en tableau 2 colonnes (logo placeholder | titre | meta) ──
    head = doc.add_table(rows=1, cols=3)
    head.autofit = False
    _set_col_widths(head, [3.0, 11.0, 4.0])
    c_logo, c_title, c_meta = head.rows[0].cells

    p = c_logo.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KAYDAN")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)
    p2 = c_logo.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("REAL ESTATE")
    r2.font.size = Pt(7)
    r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    p = c_title.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Date : {meeting.scheduled_start:%-d %B %Y}")
    r.italic = True
    r.font.size = Pt(11)
    p2 = c_title.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("RELEVÉ DE CONCLUSIONS DU\nCOMITÉ DE DIRECTION")
    r2.bold = True
    r2.font.size = Pt(12)

    p = c_meta.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Diffusion :\n").bold = True
    p.add_run("Tous les membres").font.size = Pt(9)

    # ─── Section Présents / Absents ──
    presents = list(meeting.participants.select_related("user").filter(
        attendance__status=AttendanceStatus.PRESENT,
    ).distinct())
    absents = list(meeting.participants.select_related("user").filter(
        attendance__status=AttendanceStatus.ABSENT,
    ).distinct())

    info_table = doc.add_table(rows=1, cols=2)
    info_table.autofit = False
    _set_col_widths(info_table, [13.0, 5.0])
    c_main, c_side = info_table.rows[0].cells

    p = c_main.paragraphs[0]
    p.add_run(f"Présents : {len(presents)}").bold = True
    for participant in presents:
        line = c_main.add_paragraph(style="List Bullet")
        name = (
            f"{participant.user.first_name} {participant.user.last_name}"
            if participant.user else participant.external_name
        )
        run = line.add_run(name)
        run.bold = True
        if participant.user and getattr(participant.user, "title", None):
            line.add_run(f" / {participant.user.title}").font.size = Pt(9)

    if absents:
        p = c_main.add_paragraph()
        p.add_run(f"\nAbsents : {len(absents)}").bold = True
        for participant in absents:
            line = c_main.add_paragraph(style="List Bullet")
            name = (
                f"{participant.user.first_name} {participant.user.last_name}"
                if participant.user else participant.external_name
            )
            line.add_run(name).bold = True

    # Side : président + rapporteur + ref + pages
    p = c_side.paragraphs[0]
    p.add_run("Présidence du CODIR :").bold = True
    if meeting.chair:
        c_side.add_paragraph(
            f"{meeting.chair.first_name} {meeting.chair.last_name}"
        )
    c_side.add_paragraph()
    p2 = c_side.add_paragraph()
    p2.add_run("Rapporteur :").bold = True
    if meeting.secretary:
        c_side.add_paragraph(
            f"{meeting.secretary.first_name} {meeting.secretary.last_name}"
        )

    # ─── Ordre du jour ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Ordre du jour :").bold = True
    # Items d'ordre du jour si tu en as via agendas
    try:
        for item in meeting.agenda_items.all().order_by("order"):
            doc.add_paragraph(item.title, style="List Bullet")
    except Exception:  # pragma: no cover
        doc.add_paragraph("Suivi des projets par Direction", style="List Bullet")

    doc.add_paragraph()

    # ─── Tableau des décisions / actions ──
    decisions = (
        Decision.objects.filter(meeting=meeting)
        .select_related("direction", "responsible")
        .prefetch_related("action_plan__tasks__assignee")
        .order_by("direction__name", "ref")
    )

    cols = ["DIRECTIONS", "PROJET / SUJET", "ACTIONS À MENER", "RESPONSABLES", "DEADLINE", "ÉTAT", "COMMENTAIRES"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    _set_col_widths(table, [2.8, 2.4, 4.0, 2.5, 1.7, 1.4, 3.2])

    # Header
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(c)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "1E293B")  # slate-800

    # Lignes — une ligne par décision (action). Si la décision a plusieurs tasks,
    # on liste les responsables séparés par virgule.
    current_direction = None
    for dec in decisions:
        plan = getattr(dec, "action_plan", None)
        tasks = list(plan.tasks.all()) if plan else []

        dir_name = dec.direction.name if dec.direction else ""
        same_dir = (dir_name == current_direction)
        current_direction = dir_name

        cells = table.add_row().cells

        cells[0].paragraphs[0].add_run("" if same_dir else dir_name).font.size = Pt(8)
        cells[1].paragraphs[0].add_run(dec.title).font.size = Pt(8)
        cells[2].paragraphs[0].add_run(dec.description_md or "").font.size = Pt(8)

        # Responsables
        assignees = []
        if dec.responsible:
            assignees.append(f"{dec.responsible.first_name} {dec.responsible.last_name}".strip())
        for t in tasks:
            if t.assignee:
                name = f"{t.assignee.first_name} {t.assignee.last_name}".strip()
                if name and name not in assignees:
                    assignees.append(name)
        cells[3].paragraphs[0].add_run("\n".join(assignees) or "—").font.size = Pt(8)

        cells[4].paragraphs[0].add_run(
            dec.deadline.strftime("%d/%m/%Y") if dec.deadline else "—",
        ).font.size = Pt(8)

        # Statut : on prend le plus avancé (in_progress > todo) ou done si toutes done
        status_label = _decision_status_to_label(dec, tasks)
        run_status = cells[5].paragraphs[0].add_run(status_label)
        run_status.font.size = Pt(8)
        if status_label == "En retard":
            run_status.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        elif status_label == "En cours":
            run_status.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

        # Commentaires : on concatène les derniers commentaires de tasks
        comments = []
        for t in tasks:
            last_comment = t.comments.order_by("-created_at").first()
            if last_comment:
                comments.append(last_comment.body_md[:200])
        cells[6].paragraphs[0].add_run("\n".join(comments) or "").font.size = Pt(8)

    # Footer pied de page : référence CR
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(
        f"CODIR-{meeting.scheduled_start:%d-%m-%Y} — Généré par CODIR Executive",
    )
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _decision_status_to_label(decision, tasks) -> str:
    """Détermine le label d'état affiché dans le tableau (mirror PDF Kaydan)."""
    if not tasks:
        return _map_decision_status(decision.status)
    if all(t.status == "done" for t in tasks):
        return "Terminé"
    if any(t.status == "overdue" for t in tasks):
        return "En retard"
    if any(t.status == "in_progress" for t in tasks):
        return "En cours"
    if any(t.status == "blocked" for t in tasks):
        return "En attente"
    return "Non démarré"


def _map_decision_status(status: str) -> str:
    return {
        "proposed": "Non démarré",
        "approved": "Non démarré",
        "in_progress": "En cours",
        "completed": "Terminé",
        "postponed": "En attente",
        "cancelled": "Annulé",
    }.get(status, "Non démarré")
