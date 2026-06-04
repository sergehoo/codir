"""Export du compte rendu IA d'une réunion vers Word (.docx) et PDF.

Source : `recording.ai_minutes` (Markdown structuré généré par Claude/DeepSeek)
ou éditée manuellement par l'utilisateur via PATCH /recordings/{id}/minutes/.

Stratégie :
- DOCX : python-docx + parser Markdown maison → styles Word natifs (Heading 1,
  Heading 2, List Bullet, etc.). Lecture/édition parfaite dans Word.
- PDF  : WeasyPrint qui rend un template HTML stylé en CSS print → mise en
  page propre, header Kaydan, footer paginé. WeasyPrint est lourd côté deps
  système (Pango, Cairo) — assumé pour la qualité de rendu.

Les deux fonctions retournent `bytes` ready-to-stream via HttpResponse.
"""
from __future__ import annotations

import io
import logging
import re
from typing import Iterable

from django.conf import settings
from django.template.loader import render_to_string

logger = logging.getLogger(__name__)


# ─── Parser Markdown léger (commun docx + pdf preprocessing) ───

_HEADING_RE = re.compile(r"^(#{1,4})\s+(.+)$")
_LIST_RE = re.compile(r"^\s*[-*]\s+(.+)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)([^*]+?)(?<!\*)\*(?!\*)")
_INLINE_CODE_RE = re.compile(r"`([^`]+?)`")


def _parse_markdown_blocks(md: str) -> list[dict]:
    """Convertit Markdown en blocs structurés (sans dépendance externe).

    Retourne une liste de dicts :
      {"type": "heading", "level": 1-4, "text": "..."}
      {"type": "paragraph", "text": "..."}
      {"type": "list_item", "text": "..."}
      {"type": "blank"}

    Suffit largement pour les CR de réunion. Pas de support :
    tables, images, blocs code multilignes, blockquotes complexes.
    """
    if not md:
        return []
    blocks: list[dict] = []
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            blocks.append({"type": "blank"})
            continue
        hm = _HEADING_RE.match(line)
        if hm:
            blocks.append({
                "type": "heading",
                "level": len(hm.group(1)),
                "text": hm.group(2).strip(),
            })
            continue
        lm = _LIST_RE.match(line)
        if lm:
            blocks.append({"type": "list_item", "text": lm.group(1).strip()})
            continue
        blocks.append({"type": "paragraph", "text": line.strip()})
    return blocks


def _strip_md_inline(text: str) -> str:
    """Retire les marqueurs inline Markdown (gras, italique, code) pour DOCX
    où on traite séparément les runs avec formatage."""
    return text  # on garde le brut, parsing détaillé dans _add_inline_runs


# ─── DOCX (python-docx) ────────────────────────────────────────

def generate_minutes_docx(recording) -> bytes:
    """Produit un fichier .docx (bytes) à partir de `recording.ai_minutes`.

    Style : marges 2 cm, police Calibri 11, headings dégradés cuivre Kaydan,
    listes à puces propres. Lecture parfaite dans Word/LibreOffice.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise RuntimeError(
            "python-docx requis pour l'export DOCX. "
            "Ajouter `python-docx` dans requirements.txt"
        ) from exc

    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── Header doc : titre + meta ──
    meeting = recording.meeting
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("COMPTE RENDU CODIR")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x1A, 0x16, 0x14)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_p = doc.add_paragraph()
    subtitle_run = subtitle_p.add_run(meeting.title or "Réunion CODIR")
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(0xB8, 0x69, 0x3C)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if meeting.scheduled_start:
        date_p = doc.add_paragraph()
        date_run = date_p.add_run(
            meeting.scheduled_start.strftime("%A %d %B %Y à %H:%M")
        )
        date_run.italic = True
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(0x6B, 0x5D, 0x4F)
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Trait séparateur (paragraphe vide avec bordure bas)
    doc.add_paragraph()

    # ── Corps Markdown → DOCX ──
    md = recording.ai_minutes or recording.summary or ""
    blocks = _parse_markdown_blocks(md)
    for block in blocks:
        t = block["type"]
        if t == "blank":
            continue
        if t == "heading":
            level = block["level"]
            # Mapping niveau Markdown → style DOCX
            style_name = {1: "Heading 1", 2: "Heading 2",
                          3: "Heading 3", 4: "Heading 4"}.get(level, "Heading 4")
            p = doc.add_paragraph(style=style_name)
            run = p.add_run(block["text"])
            run.font.size = Pt({1: 18, 2: 15, 3: 13, 4: 12}.get(level, 11))
            run.font.color.rgb = RGBColor(0xB8, 0x69, 0x3C) if level <= 2 \
                else RGBColor(0x3A, 0x30, 0x2A)
        elif t == "list_item":
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, block["text"])
        else:  # paragraph
            p = doc.add_paragraph()
            _add_inline_runs(p, block["text"])

    # ── Footer doc (sources, mention IA) ──
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run(
        "Document généré automatiquement par l'assistant CODIR — "
        "Validation humaine recommandée avant diffusion."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xA8, 0x9E, 0x8E)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sérialise
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_inline_runs(paragraph, text: str):
    """Ajoute le texte dans le paragraphe en parsant gras (**...**) et italique (*...*)."""
    if not text:
        return
    # Pattern combiné : **gras** | *italique* | `code` | texte normal
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(tok)


# ─── PDF (WeasyPrint) ──────────────────────────────────────────

def generate_minutes_pdf(recording) -> bytes:
    """Produit un fichier .pdf (bytes) via WeasyPrint à partir du HTML rendu.

    Le template `meeting_recordings/exports/minutes.html` définit la mise en
    page (header, footer, CSS @page). Le Markdown est converti en HTML via la
    lib `markdown` (avec quelques extensions sûres).

    ⚠️ Requiert WeasyPrint + dépendances système : libpango1.0-0, libcairo2,
    libgdk-pixbuf-2.0-0, libffi-dev. Voir Dockerfile.
    """
    try:
        from weasyprint import HTML, CSS
    except ImportError as exc:
        raise RuntimeError(
            "WeasyPrint requis pour l'export PDF. "
            "Voir Dockerfile pour les deps système (Pango/Cairo)."
        ) from exc

    # Markdown → HTML
    md = recording.ai_minutes or recording.summary or ""
    html_body = _markdown_to_html(md)

    # Render template Django
    ctx = {
        "recording": recording,
        "meeting": recording.meeting,
        "html_body": html_body,
        "site_name": getattr(settings, "DEFAULT_SITE_NAME", "CODIR Executive"),
        "generated_at": _now_french(),
    }
    html_string = render_to_string(
        "meeting_recordings/exports/minutes.html", ctx,
    )

    # WeasyPrint → PDF bytes
    pdf_buf = io.BytesIO()
    HTML(string=html_string, base_url=settings.STATIC_URL or "/").write_pdf(
        pdf_buf,
    )
    return pdf_buf.getvalue()


def _markdown_to_html(md: str) -> str:
    """Convertit Markdown en HTML sécurisé.

    Utilise `markdown` lib si dispo (rendu complet). Sinon fallback parser
    maison qui couvre les cas du CR (heading, list, paragraph, bold, italic).
    """
    if not md:
        return ""
    try:
        import markdown as md_lib
        return md_lib.markdown(
            md, extensions=["extra", "sane_lists", "smarty"], output_format="html5",
        )
    except ImportError:
        return _markdown_to_html_fallback(md)


def _markdown_to_html_fallback(md: str) -> str:
    """Parser maison sans dépendance. Suffit pour les CR exécutifs."""
    blocks = _parse_markdown_blocks(md)
    out: list[str] = []
    in_list = False
    for b in blocks:
        t = b["type"]
        if t == "list_item":
            if not in_list:
                out.append("<ul>")
                in_list = True
            out.append(f"<li>{_inline_md_to_html(b['text'])}</li>")
        else:
            if in_list:
                out.append("</ul>")
                in_list = False
            if t == "heading":
                lvl = b["level"]
                out.append(f"<h{lvl}>{_inline_md_to_html(b['text'])}</h{lvl}>")
            elif t == "paragraph":
                out.append(f"<p>{_inline_md_to_html(b['text'])}</p>")
            # blank ignored
    if in_list:
        out.append("</ul>")
    return "\n".join(out)


def _inline_md_to_html(text: str) -> str:
    """Convertit gras / italique / code inline."""
    # Escape HTML basique
    escaped = (text.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;"))
    escaped = _BOLD_RE.sub(r"<strong>\1</strong>", escaped)
    escaped = _INLINE_CODE_RE.sub(r"<code>\1</code>", escaped)
    # Italique seulement après gras/code pour ne pas matcher dans **
    escaped = _ITALIC_RE.sub(r"<em>\1</em>", escaped)
    return escaped


def _now_french() -> str:
    """Date+heure en format FR pour le footer."""
    from datetime import datetime
    return datetime.now().strftime("%d/%m/%Y à %H:%M")
